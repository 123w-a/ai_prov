# -*- coding: utf-8 -*-
"""按 AIC 第八届算法创新赛·赛题4「AI+场景创新」官方模板骨架生成参赛方案 docx。
真实内容与工程均已落地；需用户实地采集的数据（问卷/访谈/试点）一律标【待填】。
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "宋体"
OUT = r"C:/Users/31834/Desktop/小膳管家_参赛方案_算法创新赛.docx"


def set_cjk(run, font=FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def title(doc, text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    set_cjk(r)
    return p


def heading(doc, text, level):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes[level])
    set_cjk(r)
    return p


def body(doc, text, size=12, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    set_cjk(r)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)
    set_cjk(r)
    return p


def numbered(doc, text, size=12):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(size)
    set_cjk(r)
    return p


def table(doc, headers, rows, size=10.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(size)
        set_cjk(run)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(size)
            set_cjk(run)
    return t


def draw_architecture(save_path):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    import matplotlib.font_manager as fm

    # 注册系统中文字体（SimHei 黑体），避免中文成豆腐块
    cjk_font = "C:/Windows/Fonts/simhei.ttf"
    if os.path.exists(cjk_font):
        fm.fontManager.addfont(cjk_font)
        cjk_name = fm.FontProperties(fname=cjk_font).get_name()
        plt.rcParams["font.sans-serif"] = [cjk_name]
        plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(11, 8.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")

    def box(x, y, w, h, text, fc, tc="white", fs=11):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                                    linewidth=1.2, edgecolor="#0C447C", facecolor=fc))
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                color=tc, fontsize=fs, fontweight="bold")

    for y, label in [(7.4, "应用层（交互入口）"), (4.0, "算法层（决策中枢 · 本作品核心）"), (0.6, "数据层（知识与画像底座）")]:
        ax.add_patch(FancyBboxPatch((0.2, y), 9.6, 2.5, boxstyle="round,pad=0.02",
                                    linewidth=1.5, edgecolor="#185FA5", facecolor="#EAF3FB"))
        ax.text(0.4, y + 2.35, label, ha="left", va="center", fontsize=12,
                fontweight="bold", color="#0C447C")

    box(0.7, 7.7, 2.85, 1.5, "在家：拍照识别冰箱\n（零输入多模态）", "#3B6D11")
    box(3.75, 7.7, 2.85, 1.5, "在外：定位/商圈\n（附近餐厅导航）", "#3B6D11")
    box(6.8, 7.7, 2.9, 1.5, "文字/语音：健康问答\n与模糊需求澄清", "#3B6D11")

    box(0.7, 4.3, 3.0, 1.6, "chef_think\nLangGraph 循环 Agent\n（思考+工具选择+反思）", "#185FA5")
    box(3.9, 4.3, 2.8, 1.6, "工具集 6 个\nRAG检索/联网搜/附近餐\n热量/运动/文件沙箱", "#185FA5")
    box(6.9, 4.3, 2.8, 1.6, "硬护栏\nnutrition_rules\n+ verify_answer 审核", "#A32D2D", tc="white")
    box(0.7, 6.05, 9.0, 0.55, "三层健康记忆 + 自适应追问（AgentMental 范式）", "#0F6E56", fs=10)

    box(0.7, 0.95, 4.4, 1.7, "权威护栏 RAG（Chroma）\n11 份国家食养指南/国标 PDF\n402 chunk + Contextual 前缀\nparser_router→ingest 入库", "#5F5E5A")
    box(5.3, 0.95, 4.4, 1.7, "动态发现 + 静态底表\nweb_search / nearby_food(mock)\n5_层 营养底表 CSV(1347 食物)\n用户健康画像档案", "#5F5E5A")

    ax.add_patch(FancyArrowPatch((5, 7.4), (5, 6.6), arrowstyle="-|>", mutation_scale=16, color="#185FA5"))
    ax.add_patch(FancyArrowPatch((5, 4.0), (5, 2.65), arrowstyle="-|>", mutation_scale=16, color="#185FA5"))
    ax.text(5.2, 5.0, "决策调用", fontsize=9, color="#185FA5")
    ax.text(5.2, 3.3, "知识/数据供给", fontsize=9, color="#185FA5")

    fig.tight_layout()
    fig.savefig(save_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(3.0)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)

    for _ in range(3):
        doc.add_paragraph()
    title(doc, "2026 第八届全球校园人工智能算法精英大赛", 18)
    title(doc, "算法创新赛 · 技术报告", 16)
    doc.add_paragraph()
    title(doc, "小膳管家", 30)
    title(doc, "——带健康护栏的多模态膳食决策 Agent", 14)
    for _ in range(4):
        doc.add_paragraph()
    body(doc, "团队名称：____________________    团队编号：__________", align=WD_ALIGN_PARAGRAPH.CENTER)
    body(doc, "作品名称：小膳管家（AI+场景创新 · 赛题四）", align=WD_ALIGN_PARAGRAPH.CENTER)
    body(doc, "提交日期：2026 年 ____ 月", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    heading(doc, "作品简介", 1)
    body(doc,
         "小膳管家是一个带健康护栏的多模态膳食决策 Agent，面向“不会做饭、预算紧、有慢病忌口，"
         "却每天必须决定吃什么”的人群。用户拍一张冰箱照或口述身体状态，系统在过敏/慢病忌口、"
         "控糖控盐、热量目标等约束内，给出安全、便宜、可执行的一餐方案——在家给菜谱，在外给点单红线。"
         "本项目以 11 份国家权威食养指南与国标构建可溯源 RAG 知识护栏，用 LangGraph 循环 Agent 做决策中枢，"
         "并设确定性硬规则校验节点（nutrition_rules + verify_answer），对违禁推荐带次数上限地打回重生成。"
         "创新不在“做菜技巧”，而在把国家膳食指南变成可实时调用的决策护栏，替用户完成“能不能吃、怎么搭配、去哪吃”的决策。"
         "（本段约 290 字，提交前可再精简至 300 字内）")
    doc.add_page_break()

    heading(doc, "一、项目概述", 1)
    heading(doc, "（一）背景与意义", 2)
    heading(doc, "1. 行业与民生背景", 3)
    body(doc,
         "“每天吃什么”是最高频的家庭决策痛点，但在三重趋势下被显著放大：其一，外卖成为城市青年与双职工家庭的常态，"
         "却普遍高油高盐、营养结构失衡；其二，高血压、糖尿病、痛风等慢病与超重肥胖明显年轻化，"
         "《中国居民膳食指南（2022）》《健康中国行动（2019—2030）》《“体重管理年”活动实施方案（2024）》"
         "均把“合理膳食”列为国民健康第一道防线；其三，独居、适老家庭缺乏“会做饭且懂健康”的家庭成员，"
         "决策疲劳与“吃错”风险叠加。现有工具（菜谱 App、营养计数 App、通用聊天机器人）要么只做信息搬运、"
         "要么缺乏权威健康约束，无法替用户完成“安全+便宜+可执行”的一餐决策。")
    heading(doc, "2. 目标定位", 3)
    body(doc,
         "小膳管家 = 带健康护栏的 AI 膳食决策中枢。它服务的不是“想学做菜的人”，而是“不会做饭、预算紧、"
         "还有慢病忌口，但每天必须决定吃什么的人”。输入“我有什么食材 / 我在哪 / 我有什么禁忌”，"
         "输出一顿安全、便宜、好做的一餐——在家给菜谱，在外给点单红线。核心定位是把国家权威膳食指南"
         "变成可实时调用的决策护栏，而非把上千个菜谱丢给用户自己筛。")

    heading(doc, "（二）赛题方向定位", 2)
    body(doc,
         "本作品参加算法创新赛·赛题四「AI+场景创新」。选择依据：① 场景新颖——将“菜谱搜索器”重构为"
         "“带健康护栏的膳食决策 Agent”，突破“你说啥我做啥”的被动模式；② 技术适配——多模态输入、"
         "循环 Agent、可溯源 RAG、确定性硬护栏与场景需求高度耦合，技术是“赋能场景”而非点缀；"
         "③ 落地可行——底层引擎已真实跑通，并设计了“工具→服务→平台”的演进路线（下图），"
         "诚实区分“已做 / 近期 / 远期”，避免 oversell。")
    bullet(doc, "① 教人做（已落地）：AI 菜谱+步骤，用户自己动手，规模轻——竞品多，需钩子。")
    bullet(doc, "② 附近符合身体的食物（近期可做）：定位+健康画像，从标签库筛低钠/低嘌呤等合规选项。")
    bullet(doc, "③ 上门 O2O 平台（远期愿景）：连接家庭↔持证厨师，解决“懒得做又嫌外卖没营养”。"
                "规模巨大，需团队/资本与线下合规，本作品仅做 AI 决策大脑与演示闭环，不自建平台。")
    body(doc, "共享 AI 大脑（本作品核心·竞赛主秀）：营养规划 · 健康护栏 · 需求解析 · 匹配逻辑 · "
              "菜单生成 · 溯源透明 · 三层健康记忆。", bold=True)
    doc.add_page_break()

    heading(doc, "二、主内容（解决方案 · 实施 · 验证 · 应用 · 展望）", 1)

    heading(doc, "（一）需求分析", 2)
    heading(doc, "1. 问题剖析", 3)
    body(doc, "普通菜谱类产品的三重瓶颈，恰好对应“场景创新”要解决的真痛点：")
    table(doc,
          ["瓶颈", "用户真实感受", "为什么难破"],
          [["信息搬运", "搜“番茄鸡蛋”出 1000 个结果，自己筛", "只做检索，没做决策"],
           ["无健康意识", "“我有高血压，这道菜能不能吃？”App 不知道", "缺少权威医学/营养知识库"],
           ["决策疲劳", "每天最高频的痛就是“早中晚吃什么”", "没有针对用户画像的个性化推荐"]])
    body(doc, "本作品的关键解法：把“权威知识”作为硬护栏接入决策，让每次推荐都经过"
              "“能不能吃→怎么便宜做→附近有没有”三层决策，而非仅检索。")
    heading(doc, "2. 用户需求调研", 3)
    body(doc, "【调研方法】采用便利抽样 + 社区/校园定点，目标问卷 80–100 份、深度访谈 8–10 人，"
              "人群覆盖老年人/慢病家属/独居青年各约 1/3；工具为线上问卷 + 线下纸质；"
              "局限性（诚实披露）：老年样本地域偏差、学生样本偏多等。【待填：真实抽样方式与偏差说明】")
    body(doc, "问卷核心维度（已设计，可直接发卷）：①基本信息（年龄、慢病/禁忌）；②痛点程度（1–5 分）；"
              "③使用习惯（是否用过菜谱 App、是否愿拍冰箱照）；④信任与隐私（AI 建议信任度、是否愿告知慢病信息）；"
              "⑤功能优先级排序（拍照识菜/控糖控盐/省钱快手/适老大字语音/边角料不浪费）。")
    body(doc, "【待填：回收后填入下表真实统计】", bold=True)
    table(doc,
          ["指标", "结果", "备注"],
          [["有效问卷数", "【待填】", ""],
           ["年龄分布（>60 占比）", "【待填 %】", ""],
           ["痛点 Top3", "【待填】", "按选择频次"],
           ["功能需求排序", "【待填】", "按均值"],
           ["AI 建议信任度均值", "【待填 /5】", ""],
           ["愿告知慢病信息占比", "【待填 %】", "隐私顾虑证据"]])
    body(doc, "结论模板（基于数据后填写）：调研显示【待填，如“73% 受访者因健康顾虑对日常饮食焦虑，"
              "但仅 21% 信任现有 App 建议”】，印证本作品以“健康护栏+透明标注”切入的必要性。")

    heading(doc, "（二）解决方案设计", 2)
    heading(doc, "1. 技术路线规划", 3)
    body(doc, "整体技术路线：多模态输入 → LangGraph 循环 Agent 决策中枢 → 双支柱（权威护栏 RAG + 动态发现）→ "
              "结构化输出；并在中枢后接确定性硬护栏审核。具体选型：")
    bullet(doc, "多模态输入：拍照识冰箱（视觉模型）、定位/商圈、文字/语音健康状态；")
    bullet(doc, "决策中枢：LangGraph 状态图，节点含 condense_history / chef_think / run_tools / structure_answer，"
                "工具结果回流 LLM 形成“思考→调工具→再思考”闭环；")
    bullet(doc, "权威护栏 RAG：11 份国家食养指南与国标 PDF → parser_router 确定性解析路由 → ingest 切片 + "
                "Contextual 前缀 → Chroma 向量库（402 chunk）；")
    bullet(doc, "动态发现：web_search（菜谱/成品图）、nearby_food（mock 兜底，真实接高德/百度 POI）；")
    bullet(doc, "硬护栏：nutrition_rules（确定性禁忌规则，每条带 PDF 出处）+ verify_answer（命中违禁带反馈打回重生成，"
                "最多 3 次防循环，仍不过则降级并附安全警示，绝不静默放行）；")
    bullet(doc, "记忆与交互：三层健康记忆（用户/主题/陈述）+ 自适应追问澄清（AgentMental 范式），解决表达含糊死穴；")
    bullet(doc, "透明标注：AI 生成图片强制标“AI 生成示意图”，健康建议强制附出处文件名。")

    heading(doc, "2. 系统架构设计", 3)
    body(doc, "系统按“数据层—算法层—应用层”三层解耦，确保可扩展性与稳定性：")
    arch_png = os.path.join(os.path.dirname(OUT), "小膳管家_系统架构图.png")
    try:
        draw_architecture(arch_png)
        doc.add_picture(arch_png, width=Cm(16.5))
        cap = doc.paragraphs[-1]
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        body(doc, "图 1 系统架构图（数据流：应用层入口 → 算法层决策中枢 → 数据层知识与画像供给）",
             size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    except Exception as e:
        body(doc, "【建议此处插入系统架构图：应用层（拍照/定位/问答）→ 算法层（chef_think 循环 Agent + "
                  "6 工具 + 硬护栏 + 三层记忆）→ 数据层（权威 RAG / 动态发现 / 营养底表 / 用户画像）。"
                  "绘图失败原因：" + str(e)[:80] + "】")
    body(doc, "数据层：权威护栏 RAG（Chroma，402 chunk + Contextual）、动态发现（mock）、5_层营养底表 CSV"
              "（1347 种食物）、用户健康画像档案。算法层：LangGraph 决策中枢 + 6 个工具 + 确定性硬护栏"
              " + 三层记忆。应用层：在家拍照、在外定位、文字/语音问答三类入口统一收口到同一决策中枢。")

    heading(doc, "3. 功能模块设计", 3)
    table(doc,
          ["模块", "功能", "实现方式", "协同"],
          [["知识库构建", "权威 PDF→可检索向量", "parser_router 路由 + ingest 切片 + Contextual + Chroma", "喂入 RAG 工具"],
           ["决策中枢", "多轮思考+工具编排", "LangGraph 状态图 + 循环边 + checkpointer", "调用全部工具"],
           ["工具集(6)", "RAG/联网/附近餐/热量/运动/文件", "@tool 装饰器 + 文件沙箱白名单", "被中枢调用"],
           ["硬护栏", "违禁拦截+重生成", "nutrition_rules 规则 + verify_answer 节点", "阻断不合规输出"],
           ["记忆/追问", "跨轮画像+澄清", "三层记忆 + AgentMental 范式", "提升决策个性化"]])

    heading(doc, "（三）项目实施", 2)
    heading(doc, "1. 实施计划", 3)
    table(doc,
          ["阶段", "任务", "周期", "依赖"],
          [["1", "品牌/文案/UI 全量替换为「小膳管家」", "第1周", "—"],
           ["2", "完善 nearby_food 真实地图 API 接入 + mock 兜底", "第2–3周", "高德/百度 key"],
           ["3", "拍照识别降级方案（手动勾选/文字/小票）", "第2周", "—"],
           ["4", "RAG 外延：忌口问答/带量食谱/辟谣", "第3–4周", "指南已入库"],
           ["5", "菜系级护栏 playbook（火锅/烧烤/奶茶红线）", "第4周", "从指南派生"],
           ["6", "社区/校园小范围试点，回收完成率与满意度", "第5–6周", "—"],
           ["7", "录演示视频、整理验证数据、写合规说明", "第7周", "—"]])
    heading(doc, "2. 技术实现", 3)
    body(doc, "关键技术实现细节（已落地，可复现）：")
    bullet(doc, "PDF 解析路由：基于“扫描件判定（每页字符数）/页数/图文密度/表格密度”四项加权确定性打分，"
                ">2.0 走云端 OCR 解析器，否则走本地免费解析器；纯结构属性不调用 LLM，降本且可解释。")
    bullet(doc, "离线索引：Markdown 按标题切片、PDF 按段落递归切片；Contextual Retrieval（Anthropic 2024）"
                "为每个 chunk 生成“在全文中的位置”前缀，召回精度提升实测（见验证章）。")
    bullet(doc, "循环 Agent：run_tools 执行后回流 chef_think 再次思考；长对话用 RemoveMessage 真实删旧消息，"
                "并成对清理孤儿 ToolMessage 防 400 错误；SQLite checkpointer 实现多轮记忆。")
    bullet(doc, "硬护栏：verify_answer 节点置于 chef_think 与 structure_answer 之间，命中 nutrition_rules 违禁即"
                "带“哪道菜踩了哪条禁忌+出处”反馈打回，最多 3 次后降级并附警示。")
    heading(doc, "3. 团队协作", 3)
    body(doc, "【待填：成员专业背景、分工与协作机制。例如：算法/工程（Agent 与 RAG）、"
              '产品设计（适老化与交互）、调研（问卷与试点）等，并说明跨专业协作方式。】')

    heading(doc, "（四）测试与验证", 2)
    heading(doc, "1. 测试方案", 3)
    body(doc, "设计功能/适老/可靠/合规四类测试，并设 A/B 对照实验（见下）：")
    table(doc,
          ["用例", "场景", "预期", "实际", "结果"],
          [["F-01", "痛风用户 + 海鲜菜单", "拦截并警示", "【待填】", "【待填】"],
           ["F-02", "拍冰箱照", "识别食材并出餐", "【待填】", "【待填】"],
           ["U-01", "大字模式可读性", "60 岁可读", "【待填】", "【待填】"],
           ["U-02", "语音交互完成率", "≥80%", "【待填】", "【待填】"],
           ["R-01", "含糊输入“我不舒服吃啥”", "触发追问澄清", "【待填】", "【待填】"],
           ["C-01", "AIGC 图标注", "强制“AI生成示意图”", "【待填】", "【待填】"]])
    body(doc, "对照实验：A 组（无 Agent，自行查菜谱+手动算营养）vs B 组（小膳管家），指标为任务完成率、"
              "满意度(1–5)、决策耗时(分钟)、健康合规率。每组建档 10–15 人，周期 2–4 周。【待填真实数字】")
    heading(doc, "2. 验证结果", 3)
    body(doc, "已实测的引擎级结果（无需 LLM key 即可复现）：")
    bullet(doc, "RAG 召回：Contextual 重建后，「高血压饮食」查询距离 0.178→0.171、「痛风忌口」0.280→0.263，"
                "碎片片段借助上下文前缀被正确召回；402 chunk 全部带 contextual 标记，分层命中正确。")
    bullet(doc, "硬护栏：痛风坏菜单（老火汤炖猪肝+啤酒）被确定性规则命中 5 条禁忌并打回；合规菜单放行；"
                "连续违禁 3 次后正确降级并附安全警示，无无限循环。")
    body(doc, "【待填：A/B 对照实验的四项指标真实数值、用户评价原话（至少 3–5 句，含年龄/身份标注）、"
              "问题分析与改进措施表。】", bold=True)

    heading(doc, "（五）应用效果与成果", 2)
    heading(doc, "1. 实际应用情况", 3)
    body(doc, "【待填：试点范围、对象、方式与实际案例。例如“在 X 社区/校园邀请 N 人试用 2 周，"
              '完成率 X%、典型反馈：……”。应用效果（15 分）主要来自真实落地证据，务必补。】')
    heading(doc, "2. 成果展示", 3)
    bullet(doc, "技术成果：可运行的小膳管家原型（LangGraph 循环 Agent + RAG 402 chunk + 硬护栏节点），"
                "源码可复现（见附录）。")
    bullet(doc, "知识资产：11 份国家权威食养指南/国标已结构化入库并逐条可溯源。")
    bullet(doc, "差异化亮点：健康约束硬执行（确定性规则兜底，不靠 LLM 自觉）、全链路透明标注、三层健康记忆。")
    body(doc, "【待填：论文/专利/检测报告/用户使用证明等佐证材料的取得情况。】", bold=True)

    heading(doc, "（六）总结与展望", 2)
    heading(doc, "1. 成果总结", 3)
    body(doc, "本作品把“菜谱搜索器”升级为“带健康护栏的膳食决策 Agent”，核心创新在于：①场景挖掘独特性"
              "（健康约束硬执行下的主动决策）；②技术应用创新性（AgentMental 范式迁移到膳食场景）；"
              "③用户体验优化性（全链路透明标注 + 硬护栏超限降级）。底层 LangGraph Agent / 分层 RAG / "
              "确定性护栏均已真实落地，验证了“技术赋能场景”的核心逻辑。")
    heading(doc, "2. 未来展望", 3)
    body(doc, "按“工具→服务→平台”路线演进：近期补 nearby_food 真实 POI 与菜系级护栏 playbook；"
              "远期探索上门私厨 O2O 平台——但诚实说明其规模巨大（厨师招募审核、实时调度、支付托管、"
              "食安责任险、线下合规），需团队与资本，本作品定位为该平台的 AI 决策大脑，而非自建平台运营方。"
              "技术上升级长上下文 embedding 以支持 late chunking、扩充孕期等更多人群规则源（当前待补权威 PDF）。")

    doc.add_page_break()
    heading(doc, "三、附录", 1)
    heading(doc, "1. 代码与模型", 2)
    bullet(doc, "agent_graph.py：LangGraph 决策中枢（condense_history / chef_think / run_tools / structure_answer / verify_answer）。")
    bullet(doc, "nutrition_rules.py：确定性硬禁忌规则引擎（高血压/痛风/糖尿病/肾病/高脂血症/肥胖，每条带 PDF 出处）。")
    bullet(doc, "rag/ingest.py、indexing/parser_router.py、rag/retriever.py、rag/store.py：知识库构建与检索。")
    bullet(doc, "agent_tools.py：6 个工具（web_search / get_file / nearby_food / calorie_lookup / exercise_equiv / nutrition_kb_search）。")
    bullet(doc, "agent_schemas.py：ChefAnswer 含 sources 引文字段，实现权威依据进结构化卡片。")
    bullet(doc, "embedding：bge-small-zh-v1.5；向量库：Chroma（collection=dietary_kb）。")

    heading(doc, "2. 参考文献", 2)
    numbered(doc, "成人高血压食养指南（2023 年版），国家卫生健康委.")
    numbered(doc, "成人糖尿病食养指南（2023 年版），国家卫生健康委.")
    numbered(doc, "成人高脂血症食养指南（2023 年版），国家卫生健康委.")
    numbered(doc, "成人肥胖食养指南（2024 年版），国家卫生健康委.")
    numbered(doc, "成人高尿酸血症与痛风食养指南（2024 年版），国家卫生健康委.")
    numbered(doc, "成人慢性肾脏病食养指南（2024 年版），国家卫生健康委.")
    numbered(doc, "中国居民膳食指南（2022），中国营养学会.")
    numbered(doc, "GB 28050-2025 预包装食品营养标签通则；GB 7718-2025 食品标签通则.")
    numbered(doc, "健康中国行动（2019—2030）；“体重管理年”活动实施方案（2024）.")
    numbered(doc, "Anthropic. (2024). Contextual Retrieval. 用于切片上下文前缀增强.")
    numbered(doc, "LangChain / LangGraph 官方文档：循环 Agent 与工具编排.")

    heading(doc, "3. 其他材料", 2)
    body(doc, "【待填：专利证书、检测报告、用户使用证明、试点知情同意与隐私合规说明等佐证材料，"
              "一律放入佐证材料一并上传。】")
    body(doc, "合规说明（对应作品要求#6）：营养建议均标注“仅供参考，不替代执业医师/营养师”；"
              "定位非强制、可降级为选城市/商圈；AI 生成图片强制标注；文件沙箱限制读取范围杜绝路径遍历。"
              "上门场景涉及家庭位置与健康数据，方案明确由合作持证餐饮主体承担线下审核与食安责任，"
              "平台仅做匹配决策，不碰后厨。")

    doc.save(OUT)
    print("SAVED:", OUT)
    if os.path.exists(arch_png):
        print("ARCH PNG:", arch_png)


if __name__ == "__main__":
    main()
